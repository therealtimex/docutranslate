# SPDX-FileCopyrightText: 2025 RealTimeX
# SPDX-License-Identifier: MPL-2.0
import asyncio
from dataclasses import dataclass
from io import BytesIO
from typing import Self, Literal, List, Dict, Any, Tuple

import docx
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from doctranslate.agents.segments_agent import SegmentsTranslateAgentConfig, SegmentsTranslateAgent
from doctranslate.ir.document import Document
from doctranslate.translator.ai_translator.base import AiTranslatorConfig, AiTranslator


def is_image_run(run: Run) -> bool:
    """Check if a run contains an image."""
    # w:drawing is the marker for embedded images, w:pict is the marker for VML images
    return '<w:drawing' in run.element.xml or '<w:pict' in run.element.xml


@dataclass
class DocxTranslatorConfig(AiTranslatorConfig):
    """
    Configuration class for DocxTranslator.
    """
    insert_mode: Literal["replace", "append", "prepend"] = "replace"
    separator: str = "\n"


class DocxTranslator(AiTranslator):
    """
    Translator for .docx files.
    This version is optimized to handle paragraphs with mixed text and images without losing images.
    """

    def __init__(self, config: DocxTranslatorConfig):
        super().__init__(config=config)
        self.chunk_size = config.chunk_size
        self.translate_agent = None
        if not self.skip_translate:
            agent_config = SegmentsTranslateAgentConfig(
                custom_prompt=config.custom_prompt,
                to_lang=config.to_lang,
                base_url=config.base_url,
                api_key=config.api_key,
                model_id=config.model_id,
                temperature=config.temperature,
                thinking=config.thinking,
                concurrent=config.concurrent,
                timeout=config.timeout,
                logger=self.logger,
                glossary_dict=config.glossary_dict,
                retry=config.retry,
                system_proxy_enable=config.system_proxy_enable
            )
            self.translate_agent = SegmentsTranslateAgent(agent_config)
        self.insert_mode = config.insert_mode
        self.separator = config.separator

    def _pre_translate(self, document: Document) -> Tuple[DocumentObject, List[Dict[str, Any]], List[str]]:
        """
        [Refactored] Preprocess .docx file, extract text at Run level to avoid breaking images.
        :param document: Document object containing .docx file content.
        :return: A tuple containing:
                 - docx.Document object
                 - A list containing text block information (each element represents a group of consecutive text runs)
                 - A list containing all original texts to be translated
        """
        doc = docx.Document(BytesIO(document.content))
        elements_to_translate = []
        original_texts = []

        def process_paragraph(para: Paragraph):
            nonlocal elements_to_translate, original_texts
            current_text_segment = ""
            current_runs = []

            for run in para.runs:
                if is_image_run(run):
                    # Encountered an image, treat previously accumulated text as one translation unit
                    if current_text_segment.strip():
                        elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                        original_texts.append(current_text_segment)
                    # Reset accumulator
                    current_text_segment = ""
                    current_runs = []
                else:
                    # Accumulate text run
                    current_runs.append(run)
                    current_text_segment += run.text

            # Handle the last text block at the end of paragraph
            if current_text_segment.strip():
                elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                original_texts.append(current_text_segment)

        # Traverse all paragraphs
        for para in doc.paragraphs:
            process_paragraph(para)

        # Traverse all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)

        return doc, elements_to_translate, original_texts

    def _after_translate(self, doc: DocumentObject, elements_to_translate: List[Dict[str, Any]],
                         translated_texts: List[str], original_texts: List[str]) -> bytes:
        """
        [Refactored] Write translated text back to corresponding text runs, preserving images and styles.
        """
        translation_map = dict(zip(original_texts, translated_texts))

        for i, element_info in enumerate(elements_to_translate):
            runs = element_info["runs"]
            original_text = original_texts[i]
            translated_text = translated_texts[i]

            # Determine final text based on insert mode
            if self.insert_mode == "replace":
                final_text = translated_text
            elif self.insert_mode == "append":
                final_text = original_text + self.separator + translated_text
            elif self.insert_mode == "prepend":
                final_text = translated_text + self.separator + original_text
            else:
                self.logger.error("Invalid DocxTranslatorConfig parameters")
                final_text = translated_text

            if not runs:
                continue

            # --- Core modification section ---
            # Intelligently distribute translated text to each run, preserving original formatting
            if len(runs) == 1:
                # If there's only one run, assign directly
                runs[0].text = final_text
            else:
                # Multiple runs case: distribute translated text according to original text character proportions
                original_lengths = [len(run.text) for run in runs]
                total_original_length = sum(original_lengths)

                if total_original_length == 0:
                    # If original length is 0, put all text in the first run
                    runs[0].text = final_text
                    for run in runs[1:]:
                        run.text = ""
                else:
                    # More precise proportional distribution of translated text
                    cumulative_length = 0
                    final_text_len = len(final_text)

                    for i, run in enumerate(runs):
                        if i == len(runs) - 1:
                            # Last run gets all remaining characters
                            run.text = final_text[cumulative_length:]
                        else:
                            # Calculate cumulative character count up to current run
                            cumulative_original = sum(original_lengths[:i+1])
                            target_pos = int(final_text_len * cumulative_original / total_original_length)

                            # Text for current run
                            run.text = final_text[cumulative_length:target_pos]
                            cumulative_length = target_pos
            # --- End of modification ---

        # Save modified document to BytesIO stream
        doc_output_stream = BytesIO()
        doc.save(doc_output_stream)
        return doc_output_stream.getvalue()

    def translate(self, document: Document) -> Self:
        """
        Synchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = self._pre_translate(document)
        if not original_texts:
            print("\nNo translatable text found in the document.")
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = self.glossary_agent.send_segments(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Call translation agent
        if self.translate_agent:
            translated_texts = self.translate_agent.send_segments(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts

        # Write translation results back to document
        document.content = self._after_translate(doc, elements_to_translate, translated_texts, original_texts)
        return self

    async def translate_async(self, document: Document) -> Self:
        """
        Asynchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = await asyncio.to_thread(self._pre_translate, document)
        if not original_texts:
            print("\nNo translatable text found in the document.")
            # Properly save and return in async environment
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = await self.glossary_agent.send_segments_async(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Asynchronously call translation agent
        if self.translate_agent:
            translated_texts = await self.translate_agent.send_segments_async(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts
        # Write translation results back to document
        document.content = await asyncio.to_thread(self._after_translate, doc, elements_to_translate, translated_texts,
                                                   original_texts)
        return self
